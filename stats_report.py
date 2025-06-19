import os
import time
import fitz  
import docx

DOCS_FOLDER = "documents"


def get_file_size(path):
    return os.path.getsize(path) / 1024  


def count_documents():
    return [f for f in os.listdir(DOCS_FOLDER) if f.lower().endswith(('.pdf', '.docx'))]


def simulate_sorting(documents):
    start = time.time()
    for filename in documents:
        path = os.path.join(DOCS_FOLDER, filename)
        try:
            if filename.lower().endswith('.pdf'):
                with fitz.open(path) as doc:
                    _ = doc[0].get_text()
            elif filename.lower().endswith('.docx'):
                doc = docx.Document(path)
                _ = doc.paragraphs[0].text
        except Exception as e:
            print(f"[!] Skipping {filename}: {e}")
            continue
    return time.time() - start


def simulate_search(documents, keyword="data"):
    start = time.time()
    for filename in documents:
        path = os.path.join(DOCS_FOLDER, filename)
        try:
            if filename.lower().endswith('.pdf'):
                with fitz.open(path) as doc:
                    for page in doc:
                        _ = keyword.lower() in page.get_text().lower()
            elif filename.lower().endswith('.docx'):
                doc = docx.Document(path)
                for para in doc.paragraphs:
                    _ = keyword.lower() in para.text.lower()
        except Exception as e:
            print(f"[!] Error searching {filename}: {e}")
            continue
    return time.time() - start


def simulate_classification(documents):
    start = time.time()
    for filename in documents:
        path = os.path.join(DOCS_FOLDER, filename)
        try:
            if filename.lower().endswith('.pdf'):
                with fitz.open(path) as doc:
                    _ = ''.join(page.get_text() for page in doc)
            elif filename.lower().endswith('.docx'):
                doc = docx.Document(path)
                _ = "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            print(f"[!] Error classifying {filename}: {e}")
            continue
    return time.time() - start


def generate_stats_report():
    docs = count_documents()
    file_count = len(docs)
    total_size = sum(get_file_size(os.path.join(DOCS_FOLDER, f)) for f in docs)
    sort_time = simulate_sorting(docs)
    search_time = simulate_search(docs)
    classify_time = simulate_classification(docs)

    return [
        "📊 Project Statistics:",
        f"- Total number of documents: {file_count}",
        f"- Total size: {total_size:.2f} KB",
        f"- Time to sort documents: {sort_time:.2f} sec",
        f"- Time to search documents: {search_time:.2f} sec",
        f"- Time to classify documents: {classify_time:.2f} sec"
    ]
