import os
import fitz  # PyMuPDF
import docx

DOCS_FOLDER = "documents"

# ✅ تمييز الكلمات داخل PDF
def search_pdf(path, keyword):
    results = []
    try:
        doc = fitz.open(path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if keyword.lower() in text.lower():
                # تمييز كل التطابقات
                found_instances = page.search_for(keyword, hit_max=1000)
                for inst in found_instances:
                    try:
                        page.add_highlight_annot(inst)
                    except Exception:
                        continue
                # استخراج الأسطر التي تحتوي على الكلمة
                for line in text.split('\n'):
                    if keyword.lower() in line.lower():
                        results.append((page_num, line.strip()))
        # حفظ الملف مع التمييز
        try:
            doc.save(path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
        except Exception as save_err:
            print(f"[!] Warning: Could not save highlights to '{path}': {save_err}")
        doc.close()
    except Exception as e:
        print(f"[!] Error reading {path}: {e}")
    return results

# ✅ تمييز الكلمة داخل DOCX فقرة واحدة
def highlight_word_in_docx(paragraph, keyword):
    words = paragraph.text.split()
    paragraph.clear()
    for word in words:
        run = paragraph.add_run(word + " ")
        if keyword.lower() in word.lower():
            run.font.highlight_color = 7  # Yellow

# ✅ البحث داخل ملفات DOCX
def search_docx(path, keyword):
    results = []
    try:
        doc = docx.Document(path)
        changed = False
        for para in doc.paragraphs:
            if keyword.lower() in para.text.lower():
                results.append(para.text.strip())
                highlight_word_in_docx(para, keyword)
                changed = True
        if changed:
            doc.save(path)
    except Exception as e:
        print(f"[!] Error reading {path}: {e}")
    return results

# ✅ بحث شامل في جميع الملفات داخل المجلد
def search_documents(keyword):
    result_dict = {}
    for filename in os.listdir(DOCS_FOLDER):
        full_path = os.path.join(DOCS_FOLDER, filename)
        matches = []

        if filename.lower().endswith(".pdf"):
            matches = search_pdf(full_path, keyword)
            if matches:
                result_dict[filename] = [f"Page {p}: {line}" for p, line in matches]

        elif filename.lower().endswith(".docx"):
            matches = search_docx(full_path, keyword)
            if matches:
                result_dict[filename] = matches

    return result_dict
